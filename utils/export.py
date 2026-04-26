from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
from datetime import datetime


def export_to_docx(project_info, requirements, srs_content):
    """Export SRS ke file Word (.docx)"""
    doc = Document()

    # ── Cover Page ──────────────────────────────────────────
    title = doc.add_heading('SOFTWARE REQUIREMENTS SPECIFICATION', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading(project_info.get('name', 'Sistem'), 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('')
    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = 'Table Grid'
    rows_data = [
        ('Tanggal', datetime.now().strftime('%d %B %Y')),
        ('Versi', '1.0'),
        ('Status', 'Draft'),
        ('Domain', project_info.get('domain', '-')),
    ]
    for i, (label, value) in enumerate(rows_data):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[1].text = value

    doc.add_page_break()

    # ── SRS Content ─────────────────────────────────────────
    for line in srs_content.split('\n'):
        line = line.strip()
        if not line:
            continue
        if line.startswith('# '):
            doc.add_heading(line[2:], 1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], 2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], 3)
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
        else:
            doc.add_paragraph(line)

    doc.add_page_break()

    # ── Traceability Matrix ──────────────────────────────────
    doc.add_heading('Traceability Matrix', 1)
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    headers = ['ID', 'Requirement', 'Tipe', 'Prioritas', 'Skor Kualitas']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    for req in requirements:
        row = table.add_row().cells
        row[0].text = req.get('id', '-')
        row[1].text = req.get('text', '')[:80] + ('...' if len(req.get('text', '')) > 80 else '')
        row[2].text = req.get('type', '-')
        row[3].text = req.get('priority', '-')
        row[4].text = str(req.get('quality_score', '-'))

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer